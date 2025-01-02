#Extract the transcript from YouTube video and save it to the text file.
from langchain_community.document_loaders import YoutubeLoader
from langchain_community.docstore.document import Document
import textwrap
import string 
from pathlib import Path 
from urllib.request import urlopen
from bs4 import BeautifulSoup 
import docx as dx

url = input("Paste your YouTube url: ")

soup = BeautifulSoup(urlopen(url), features="lxml")
    
webpage_title = soup.title.get_text()    
webpage_title = webpage_title.replace(" - YouTube", "")    
    
exclude = set(string.punctuation)
filename_title = "".join(ch for ch in webpage_title if ch not in exclude)

loader = YoutubeLoader.from_youtube_url(
    url    
)
docs = loader.load()
for doc in docs:  
    document = Document(
        page_content = doc.page_content        
    )
        
    content_text = textwrap.fill(document.page_content, width=80) 

    full_content = filename_title + "\n" \
        + content_text    
    
    path_filename_title = Path(filename_title + ".txt")    
    
    path_filename_title.write_text(full_content, encoding='utf-8')
    
    document = dx.Document()
    document.add_paragraph(path_filename_title.read_text(encoding='utf-8'))
    document.save(filename_title + ".docx")